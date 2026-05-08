#!/usr/bin/env python3
import os
import sys
import json
import argparse
import re
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

def add_paragraph_after(paragraph):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)

def add_table_after(paragraph, doc, rows=0, cols=2):
    tmp_table = doc.add_table(rows=rows, cols=cols)
    tbl_elem = tmp_table._tbl
    doc.element.body.remove(tbl_elem)
    paragraph._p.addnext(tbl_elem)
    return Table(tbl_elem, paragraph._parent)

def remove_table_header_formatting(table):
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), "FFFFFF")
            tcPr.append(shading)


# ============================================================
# Eliminación de secciones nulas
# ============================================================

W_P   = qn('w:p')
W_TBL = qn('w:tbl')
W_T   = qn('w:t')


def _paragraph_full_text(p_elem):
    """Texto plano concatenando todos los <w:t> del párrafo (independiente de runs)."""
    return ''.join((t.text or '') for t in p_elem.iter(W_T))


def _normalize_paragraph_runs(p_elem):
    """
    Si el párrafo contiene un marcador {{naN-X}} partido entre múltiples runs,
    consolida el texto en el primer <w:t> y vacía los demás <w:t>. Esto
    permite usar replace simple después.
    """
    full = _paragraph_full_text(p_elem)
    if '{{na' not in full:
        return
    t_elements = list(p_elem.iter(W_T))
    if not t_elements:
        return
    t_elements[0].text = full
    for t in t_elements[1:]:
        t.text = ''


def _clear_markers_in_paragraph(p_elem, markers):
    """Reemplaza cada marker por '' en los <w:t> del párrafo (tras normalizar)."""
    _normalize_paragraph_runs(p_elem)
    for t in p_elem.iter(W_T):
        if not t.text:
            continue
        for m in markers:
            if m in t.text:
                t.text = t.text.replace(m, '')


def _find_marker_paragraph(doc, marker):
    """Busca el primer <w:p> que contenga `marker` en cualquier parte del documento."""
    body = doc.element.body
    for p in body.iter(W_P):
        if marker in _paragraph_full_text(p):
            return p
    return None


def _ancestor_in_body(elem, body):
    """Sube por los ancestros hasta encontrar uno cuyo padre sea el body."""
    cur = elem
    while cur is not None:
        parent = cur.getparent()
        if parent is body:
            return cur
        cur = parent
    return None


def process_section_markers(doc, sections_to_remove, log=None):
    """
    Para cada sección N de 1 a 11:
      - Busca {{naN-1}} y {{naN-2}}.
      - Si N está en sections_to_remove: elimina del body todos los bloques
        desde el ancestro-en-body del marcador inicial hasta el del marcador
        final (inclusive).
      - Si no: solo limpia los marcadores del texto.
    """
    body = doc.element.body
    if log is None:
        log = lambda msg: None

    for n in range(1, 12):
        start_marker = "{{na%d-1}}" % n
        end_marker   = "{{na%d-2}}" % n

        start_p = _find_marker_paragraph(doc, start_marker)
        end_p   = _find_marker_paragraph(doc, end_marker)

        if start_p is None and end_p is None:
            log("Seccion %d: no se encontraron marcadores." % n)
            continue

        if start_p is None or end_p is None:
            log("Seccion %d: marcadores desbalanceados (start=%s, end=%s)."
                % (n, start_p is not None, end_p is not None))
            for p in (start_p, end_p):
                if p is not None:
                    _clear_markers_in_paragraph(p, [start_marker, end_marker])
            continue

        if n not in sections_to_remove:
            _clear_markers_in_paragraph(start_p, [start_marker])
            _clear_markers_in_paragraph(end_p, [end_marker])
            log("Seccion %d: marcadores limpiados (no se elimina tabla)." % n)
            continue

        start_block = _ancestor_in_body(start_p, body)
        end_block   = _ancestor_in_body(end_p, body)

        if start_block is None or end_block is None:
            log("Seccion %d: no se encontro ancestro en body. Solo limpio marcadores." % n)
            _clear_markers_in_paragraph(start_p, [start_marker])
            _clear_markers_in_paragraph(end_p, [end_marker])
            continue

        if start_block is end_block:
            log("Seccion %d: ambos marcadores en el mismo bloque del body. Solo limpio marcadores." % n)
            _clear_markers_in_paragraph(start_block, [start_marker, end_marker])
            continue

        to_remove = []
        cur = start_block
        while cur is not None:
            to_remove.append(cur)
            if cur is end_block:
                break
            cur = cur.getnext()

        if not to_remove or to_remove[-1] is not end_block:
            log("Seccion %d: end_block no esta despues de start_block. Solo limpio marcadores." % n)
            _clear_markers_in_paragraph(start_p, [start_marker])
            _clear_markers_in_paragraph(end_p, [end_marker])
            continue

        log("Seccion %d: eliminando %d bloque(s) del body." % (n, len(to_remove)))
        for elem in to_remove:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)


# ============================================================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--folder", required=True, help="Directorio donde estan el DOCX, mapping.json e imagenes")
    parser.add_argument("--docx",   required=True, help="Nombre del archivo DOCX dentro de --folder")
    parser.add_argument("--token",  required=True, help="Texto a buscar/reemplazar")
    parser.add_argument("--sections", required=False, default=None,
                        help="Ruta al JSON con la lista de secciones a eliminar.")
    args = parser.parse_args()

    folder_path = args.folder
    docx_name   = args.docx
    token       = args.token

    docx_path = os.path.join(folder_path, docx_name)
    mapping_path = os.path.join(folder_path, "mapping.json")
    signatures_path = os.path.join(folder_path, "signatures.json")

    doc = Document(docx_path)

    # ---- Procesar marcadores de secciones nulas ----
    sections_to_remove = []
    if args.sections and os.path.exists(args.sections):
        try:
            with open(args.sections, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    sections_to_remove = [int(x) for x in data]
        except Exception as e:
            print("Error leyendo sections JSON: {}".format(e))

    print("Secciones a eliminar: {}".format(sections_to_remove))
    process_section_markers(doc, sections_to_remove, log=lambda m: print("[secciones] " + m))
    # -----------------------------------------------

    signatures_data = {}
    if os.path.exists(signatures_path):
        with open(signatures_path, "r", encoding="utf-8") as f:
            signatures_data = json.load(f)

    def process_signatures(paragraphs):
        pattern = r'(\{\{firma_admin\}\}|\{\{firma_inspector\}\})'
        for paragraph in paragraphs:
            text = paragraph.text
            if "{{firma_admin}}" not in text and "{{firma_inspector}}" not in text:
                continue

            alignment = paragraph.alignment
            parts = re.split(pattern, text)
            paragraph.clear()

            for part in parts:
                if part == "{{firma_admin}}":
                    img_filename = signatures_data.get("{{firma_admin}}")
                    if img_filename:
                        img_path = os.path.join(folder_path, img_filename)
                        if os.path.exists(img_path):
                            try:
                                run = paragraph.add_run()
                                run.add_picture(img_path, width=Inches(1.8))
                            except Exception as e:
                                print("Error cargando firma admin: {}".format(e))

                elif part == "{{firma_inspector}}":
                    sigs = signatures_data.get("{{firma_inspector}}", [])
                    img_filename = None
                    if len(sigs) > 0:
                        img_filename = sigs.pop(0)

                    if img_filename:
                        img_path = os.path.join(folder_path, img_filename)
                        if os.path.exists(img_path):
                            try:
                                run = paragraph.add_run()
                                run.add_picture(img_path, width=Inches(1.8))
                            except Exception as e:
                                print("Error cargando firma inspector: {}".format(e))
                else:
                    if part:
                        paragraph.add_run(part)

            if alignment is not None:
                paragraph.alignment = alignment

    process_signatures(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_signatures(cell.paragraphs)

    if not os.path.exists(mapping_path):
        for paragraph in doc.paragraphs:
            if token in paragraph.text:
                paragraph.text = paragraph.text.replace(token, "")
        doc.save(docx_path)
        return

    with open(mapping_path, "r", encoding="utf-8") as f:
        photos_mapping = json.load(f)

    if not photos_mapping:
        for paragraph in doc.paragraphs:
            if token in paragraph.text:
                paragraph.text = paragraph.text.replace(token, "")
        doc.save(docx_path)
        return

    inserted = False
    for paragraph in doc.paragraphs:
        if token in paragraph.text:
            paragraph.text = paragraph.text.replace(token, "")

            new_paragraph = add_paragraph_after(paragraph)
            table = add_table_after(new_paragraph, doc, rows=0, cols=2)
            remove_table_header_formatting(table)

            for i in range(0, len(photos_mapping), 2):
                row_img  = table.add_row()
                row_text = table.add_row()

                fn1 = photos_mapping[i]["filename"]
                tx1 = photos_mapping[i]["text"]
                p_img1 = row_img.cells[0].paragraphs[0]
                p_img1.paragraph_format.keep_with_next = True
                run1 = p_img1.add_run()
                run1.add_picture(os.path.join(folder_path, fn1), width=Inches(2))

                p_txt1 = row_text.cells[0].paragraphs[0]
                p_txt1.add_run(tx1)

                if i+1 < len(photos_mapping):
                    fn2 = photos_mapping[i+1]["filename"]
                    tx2 = photos_mapping[i+1]["text"]
                    p_img2 = row_img.cells[1].paragraphs[0]
                    p_img2.paragraph_format.keep_with_next = True
                    run2 = p_img2.add_run()
                    run2.add_picture(os.path.join(folder_path, fn2), width=Inches(2))

                    p_txt2 = row_text.cells[1].paragraphs[0]
                    p_txt2.add_run(tx2)

            inserted = True
            break

    doc.save(docx_path)

if __name__ == "__main__":
    main()