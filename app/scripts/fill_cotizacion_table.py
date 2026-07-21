#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""


Uso:
    python3 fill_cotizacion_table.py <plantilla.docx> <datos.json> <salida.docx>
"""

import sys
import json
import copy

from docx import Document
from docx.table import _Row


def paragraph_text(paragraph):
    return "".join(run.text for run in paragraph.runs)


def replace_in_paragraph(paragraph, replacements):
    full_text = paragraph_text(paragraph)
    if not any(marker in full_text for marker in replacements):
        return

    new_text = full_text
    for marker, value in replacements.items():
        new_text = new_text.replace(marker, value)

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_in_cell(cell, replacements):
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, replacements)


def row_marker_score(row, markers):
    text = ""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            text += paragraph_text(paragraph)
    return sum(1 for m in markers if m in text)


def find_template_row(document, markers):
    best = None
    best_score = 0
    for table in document.tables:
        for idx, row in enumerate(table.rows):
            score = row_marker_score(row, markers)
            if score > best_score:
                best_score = score
                best = (table, idx)
    return best


def main():
    if len(sys.argv) != 4:
        sys.stderr.write(
            "Uso: fill_cotizacion_table.py <plantilla.docx> <datos.json> <salida.docx>\n"
        )
        sys.exit(2)

    template_path, data_path, output_path = sys.argv[1], sys.argv[2], sys.argv[3]

    with open(data_path, encoding="utf-8") as f:
        payload = json.load(f)

    rows = payload.get("rows", [])
    marker_map = payload.get("marker_map", {})
    markers = list(marker_map.values())

    document = Document(template_path)
    found = find_template_row(document, markers)

    if not found or not rows:
        document.save(output_path)
        return

    table, template_idx = found
    template_tr = table.rows[template_idx]._tr

    anchor = template_tr
    for record in rows:
        new_tr = copy.deepcopy(template_tr)
        anchor.addnext(new_tr)
        anchor = new_tr

        replacements = {
            marker: str(record.get(field, ""))
            for field, marker in marker_map.items()
        }
        new_row = _Row(new_tr, table)
        for cell in new_row.cells:
            replace_in_cell(cell, replacements)

    template_tr.getparent().remove(template_tr)
    document.save(output_path)


if __name__ == "__main__":
    main()